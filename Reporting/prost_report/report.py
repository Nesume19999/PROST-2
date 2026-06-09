"""
report.py
=========
Assemble a Word (.docx) report for one simulation: an auto-written narrative,
a compact indicator table, and the charts produced by charts.build_charts.
Mirrors the Stata module "R5 - Report document.do".
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .prepare import SimData

FIGURE_ORDER = [
    ("01_contributors_vs_beneficiaries.png", "Contributors vs beneficiaries"),
    ("02_dependency_ratio.png", "System dependency ratio"),
    ("03_beneficiaries_by_type.png", "Beneficiaries by pension type"),
    ("04_wage_vs_pension.png", "Average wage vs old-age pension"),
    ("05_replacement_and_coverage.png", "Replacement ratio and coverage"),
    ("06_financial_balance.png", "Illustrative financial balance"),
    ("07_age_pyramid_final_year.png", "Beneficiary age structure (final year)"),
    ("08_density_by_decile_final_year.png", "Contribution density by decile (final year)"),
]


def _val(df, year, col):
    return float(df.loc[df["year"] == year, col].iloc[0])


def build_report(sd: SimData, figdir: str, docx_path: str) -> str:
    """Build the Word report for simulation `sd`; return the path."""
    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
    sys = sd.system
    y0, y1 = sd.first_year, sd.last_year

    dep0, dep1 = _val(sys, y0, "dependency_ratio"), _val(sys, y1, "dependency_ratio")
    ben0, ben1 = _val(sys, y0, "beneficiaries_m"), _val(sys, y1, "beneficiaries_m")
    con0, con1 = _val(sys, y0, "contributors_m"), _val(sys, y1, "contributors_m")
    repl1 = _val(sys, y1, "replacement_ratio")
    neg = sys.loc[sys["balance"] < 0, "year"]
    neg_year = int(neg.min()) if len(neg) else None

    doc = Document()

    doc.add_heading("PROST v2 Pension Projection Report", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"Simulation: {sd.sim}   |   Country: {sd.params.country or 'n/a'}   "
        f"|   Generated: {date.today().isoformat()}"
    ).italic = True

    # --- 1. Summary --------------------------------------------------------
    doc.add_heading("1. Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"This report summarises the {sd.sim} projection over {y0}–{y1}. ")
    p.add_run("The system dependency ratio moves from ")
    p.add_run(f"{dep0:.1f}").bold = True
    p.add_run(" to ")
    p.add_run(f"{dep1:.1f}").bold = True
    p.add_run(f" beneficiaries per 100 contributors (a change of ")
    p.add_run(f"{dep1 - dep0:+.1f}").bold = True
    p.add_run("). The number of beneficiaries goes from ")
    p.add_run(f"{ben0:.2f}").bold = True
    p.add_run(" million to ")
    p.add_run(f"{ben1:.2f}").bold = True
    p.add_run(" million, while contributors move from ")
    p.add_run(f"{con0:.2f}").bold = True
    p.add_run(" million to ")
    p.add_run(f"{con1:.2f}").bold = True
    p.add_run(" million. The average old-age replacement ratio in ")
    p.add_run(f"{y1}").bold = True
    p.add_run(" is ")
    p.add_run(f"{repl1:.1f}%").bold = True
    p.add_run(".")
    if neg_year is not None:
        q = doc.add_paragraph()
        q.add_run(
            "On the illustrative, parameter-driven financing assumptions, the "
            "annual balance first turns negative in "
        )
        q.add_run(f"{neg_year}").bold = True
        q.add_run(".")

    # --- 2. Key indicators (decade snapshots) ------------------------------
    doc.add_heading("2. Key indicators (decade snapshots)", level=1)
    snap_years = [y for y in sys["year"]
                  if y % 10 == 0 or y in (y0, y1)]
    cols = [
        ("year", "Year", "{:.0f}"),
        ("dependency_ratio", "Depend. ratio", "{:.1f}"),
        ("coverage_active", "Coverage %", "{:.1f}"),
        ("replacement_ratio", "Repl. ratio %", "{:.1f}"),
        ("beneficiaries_m", "Benef. (m)", "{:.2f}"),
        ("contributors_m", "Contrib. (m)", "{:.2f}"),
        ("balance_b", "Balance (bn)", "{:.2f}"),
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for j, (_, header, _fmt) in enumerate(cols):
        table.rows[0].cells[j].paragraphs[0].add_run(header).bold = True
    for yr in snap_years:
        row = table.add_row().cells
        for j, (key, _h, fmt) in enumerate(cols):
            row[j].text = fmt.format(_val(sys, yr, key))

    # --- 3. Figures --------------------------------------------------------
    doc.add_heading("3. Figures", level=1)
    for fname, caption in FIGURE_ORDER:
        fpath = os.path.join(figdir, fname)
        if not os.path.exists(fpath):
            continue
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(fpath, width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

    # --- 4. Notes ----------------------------------------------------------
    doc.add_heading("4. Notes", level=1)
    doc.add_paragraph(
        "Beneficiary, contributor, wage and pension figures come directly from "
        "the PROST v2 in-year reporting outputs. Contributions, expenditure and "
        "the net balance are illustrative: they are derived from an assumed "
        f"contribution rate of {sd.params.contrib_rate}% applied to the average "
        f"wage and the number of contributors, annualised over {sd.params.periods} "
        "pay periods. Treat the financing block as indicative rather than an "
        "actuarial balance."
    )

    doc.save(docx_path)
    return docx_path
