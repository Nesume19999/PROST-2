"""
compare.py
==========
Compare two prepared simulations side by side: overlay charts, an Excel
comparison workbook, and a short Word comparison report.
Mirrors the Stata module "R6 - Scenario comparison.do".
"""

from __future__ import annotations

import os
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .prepare import SimData

C = ["#1f77b4", "#d62728"]


def _overlay(a: SimData, b: SimData, col, title, ylabel, path, zeroline=False):
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.plot(a.system["year"], a.system[col], color=C[0], lw=2, label=a.sim)
    ax.plot(b.system["year"], b.system[col], color=C[1], lw=2, ls="--", label=b.sim)
    if zeroline:
        ax.axhline(0, color="grey", lw=0.8)
    ax.set(title=title, xlabel="Year", ylabel=ylabel)
    ax.grid(alpha=0.3)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(loc="best")
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def compare_simulations(a: SimData, b: SimData, figdir: str,
                        xlsx_path: str, docx_path: str) -> dict:
    """Build comparison figures, workbook and report for two simulations."""
    os.makedirs(figdir, exist_ok=True)
    cur = a.params.currency

    # --- Overlay charts ----------------------------------------------------
    f1 = os.path.join(figdir, "C1_dependency_ratio.png")
    f2 = os.path.join(figdir, "C2_beneficiaries.png")
    f3 = os.path.join(figdir, "C3_net_balance.png")
    _overlay(a, b, "dependency_ratio",
             f"Dependency ratio: {a.sim} vs {b.sim}",
             "Beneficiaries per 100 contributors", f1)
    _overlay(a, b, "beneficiaries_m",
             f"Beneficiaries: {a.sim} vs {b.sim}",
             "Persons (millions)", f2)
    _overlay(a, b, "balance_b",
             f"Illustrative net balance: {a.sim} vs {b.sim}",
             f"{cur}, billions per year", f3, zeroline=True)

    # --- Comparison workbook ----------------------------------------------
    keep = ["year", "dependency_ratio", "beneficiaries_m", "contributors_m",
            "replacement_ratio", "balance_b"]
    da = a.system[keep].add_suffix(f"_{a.sim}").rename(columns={f"year_{a.sim}": "year"})
    db = b.system[keep].add_suffix(f"_{b.sim}").rename(columns={f"year_{b.sim}": "year"})
    comp = da.merge(db, on="year", how="outer").sort_values("year")
    for metric in ["dependency_ratio", "beneficiaries_m", "balance_b"]:
        comp[f"{metric}_diff"] = comp[f"{metric}_{b.sim}"] - comp[f"{metric}_{a.sim}"]

    os.makedirs(os.path.dirname(os.path.abspath(xlsx_path)), exist_ok=True)
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as xw:
        comp.to_excel(xw, sheet_name="Comparison", index=False)

    # --- Short comparison report ------------------------------------------
    y1 = int(comp["year"].max())
    row = comp.loc[comp["year"] == y1].iloc[0]
    dep_diff = float(row["dependency_ratio_diff"])
    bal_diff = float(row["balance_b_diff"])

    doc = Document()
    doc.add_heading("PROST v2 Scenario Comparison", level=0)
    doc.add_paragraph(f"{a.sim} vs {b.sim}   |   Generated: {date.today().isoformat()}").italic = True

    doc.add_heading("Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"By {y1}, the {b.sim} scenario changes the dependency ratio by ")
    p.add_run(f"{dep_diff:+.1f}").bold = True
    p.add_run(f" beneficiaries per 100 contributors relative to {a.sim}, and the "
              "illustrative annual balance by ")
    p.add_run(f"{bal_diff:+.2f}").bold = True
    p.add_run(f" billion {cur}.")

    doc.add_heading("Figures", level=1)
    for fpath in (f1, f2, f3):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(fpath, width=Inches(6.0))

    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
    doc.save(docx_path)

    return {"figures": [f1, f2, f3], "workbook": xlsx_path, "report": docx_path}
